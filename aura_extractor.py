#!/usr/bin/env python3
"""
AURA ERP — Extrator Universal de Bases de Fornecedores
======================================================
Converte qualquer arquivo de fornecedor para CSV padronizado
pronto para importar no AURA ERP.

Formatos suportados:
  PDF (texto digital) → pdfplumber
  PDF (escaneado)     → pdf2image + pytesseract (OCR)
  XLSX / XLS          → pandas + openpyxl/xlrd
  CSV                 → pandas
  DOCX / DOC          → python-docx
  JPG / JPEG / PNG    → Pillow + pytesseract (OCR)

Instalação das dependências:
  pip install pdfplumber pandas openpyxl xlrd python-docx pillow pytesseract pdf2image

  Windows: instalar Tesseract OCR em https://github.com/UB-Mannheim/tesseract/wiki
  Depois configurar: pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

Uso:
  python aura_extractor.py <arquivo>                    # extrai e salva CSV
  python aura_extractor.py <arquivo> --fornecedor "Nome"
  python aura_extractor.py <pasta>                      # processa todos os arquivos de uma pasta
  python aura_extractor.py <arquivo> --template         # usa/salva template de mapeamento
"""

import sys
import os
import re
import json
import argparse
import pandas as pd
from pathlib import Path

# ─── CONFIGURAÇÕES ───────────────────────────────────────────────────
TEMPLATES_FILE = "aura_templates.json"
OUTPUT_DIR = "aura_extraidos"

# Palavras-chave para detecção automática de colunas
KEYWORD_MAP = {
    "produto": ["peptideo","product","name","nome","peptide","description","desc","item","produto","compound"],
    "codigo":  ["codigo","code","sku","cod","id","ref","internal","código"],
    "dosagem": ["dosagem","dose","mg","iu","ml","unit","unidade","strength","qty"],
    "preco":   ["preco","price","usd","valor","custo","cost","preco_kit","kit10","kit","amount","rate"],
}

# ─── EXTRATORES ──────────────────────────────────────────────────────
def extract_csv(path: Path) -> pd.DataFrame:
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            for sep in [",", ";", "\t"]:
                try:
                    df = pd.read_csv(path, sep=sep, encoding=enc, dtype=str)
                    if len(df.columns) >= 2:
                        print(f"  ✓ CSV lido: {len(df)} linhas, {len(df.columns)} colunas (sep='{sep}', enc={enc})")
                        return df
                except:
                    pass
        except:
            pass
    raise ValueError("Não foi possível ler o CSV.")


def extract_xlsx(path: Path) -> pd.DataFrame:
    xl = pd.ExcelFile(path)
    print(f"  ℹ  Abas encontradas: {xl.sheet_names}")
    # Tentar cada aba e pegar a com mais dados
    best = None
    for sheet in xl.sheet_names:
        try:
            df = pd.read_excel(path, sheet_name=sheet, dtype=str)
            df.dropna(how="all", inplace=True)
            if best is None or len(df) > len(best):
                best = df
                print(f"  ✓ Aba '{sheet}': {len(df)} linhas")
        except:
            pass
    return best


def extract_pdf_text(path: Path) -> pd.DataFrame:
    try:
        import pdfplumber
    except ImportError:
        sys.exit("  ✗ Instale pdfplumber: pip install pdfplumber")

    print(f"  ℹ  Extraindo texto do PDF (método digital)...")
    rows = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            # Tentativa 1: extrair tabelas estruturadas
            tables = page.extract_tables()
            for tbl in tables:
                for row in tbl:
                    if row and any(c for c in row if c):
                        rows.append([str(c or "").strip() for c in row])

            # Tentativa 2: extrair texto linha a linha
            if not tables:
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    line = line.strip()
                    if len(line) > 3:
                        parts = re.split(r"\s{2,}|\t", line)
                        if len(parts) >= 2:
                            rows.append(parts)

    if not rows:
        raise ValueError("Nenhuma tabela encontrada no PDF. Tente com OCR (--ocr).")

    # Detectar se primeira linha é header
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    headers = rows[0] if all(not re.match(r"^\d+\.?\d*$", c) for c in rows[0] if c) else [f"col{i}" for i in range(max_cols)]
    data = rows[1:] if headers != [f"col{i}" for i in range(max_cols)] else rows
    df = pd.DataFrame(data, columns=headers[:max_cols])
    df = df[df.apply(lambda r: any(r.str.strip() != ""), axis=1)]
    print(f"  ✓ PDF (texto): {len(df)} linhas extraídas")
    return df


def extract_pdf_ocr(path: Path) -> pd.DataFrame:
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image
    except ImportError:
        sys.exit("  ✗ Instale: pip install pdf2image pytesseract pillow")

    print(f"  ℹ  OCR no PDF escaneado (pode demorar ~10-30s por página)...")
    images = convert_from_path(path, dpi=300)
    all_text = []
    for i, img in enumerate(images):
        print(f"     Página {i+1}/{len(images)}...")
        text = pytesseract.image_to_string(img, lang="por+eng", config="--psm 6")
        all_text.append(text)
    return parse_ocr_text("\n".join(all_text))


def extract_image_ocr(path: Path) -> pd.DataFrame:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        sys.exit("  ✗ Instale: pip install pytesseract pillow")

    print(f"  ℹ  OCR na imagem...")
    img = Image.open(path)
    text = pytesseract.image_to_string(img, lang="por+eng", config="--psm 6")
    return parse_ocr_text(text)


def extract_docx(path: Path) -> pd.DataFrame:
    try:
        from docx import Document
    except ImportError:
        sys.exit("  ✗ Instale python-docx: pip install python-docx")

    doc = Document(path)
    rows = []
    # Extrair tabelas do Word
    for tbl in doc.tables:
        for row in tbl.rows:
            rows.append([c.text.strip() for c in row.cells])
    if not rows:
        # Extrair parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                parts = re.split(r"\s{2,}|\t", para.text.strip())
                if len(parts) >= 2:
                    rows.append(parts)
    if not rows:
        raise ValueError("Nenhum dado encontrado no DOCX.")
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    headers = rows[0]
    df = pd.DataFrame(rows[1:], columns=headers)
    print(f"  ✓ DOCX: {len(df)} linhas extraídas")
    return df


def parse_ocr_text(text: str) -> pd.DataFrame:
    rows = []
    price_re = re.compile(r"\$?\s*(\d{1,4}(?:[.,]\d{1,2})?)\s*(?:USD|usd)?")
    for line in text.split("\n"):
        line = line.strip()
        if len(line) < 3:
            continue
        match = price_re.search(line)
        if match:
            price = match.group(1).replace(",", ".")
            product = line[:match.start()].strip()
            if product:
                rows.append({"produto": product, "preco_usd": price})
    if not rows:
        raise ValueError("OCR não encontrou preços reconhecíveis no texto.")
    df = pd.DataFrame(rows)
    print(f"  ✓ OCR: {len(df)} linhas com preços encontradas")
    return df


# ─── MAPEAMENTO DE COLUNAS ───────────────────────────────────────────
def auto_map_columns(df: pd.DataFrame) -> dict:
    mapping = {}
    cols = df.columns.tolist()
    cols_lower = [c.lower() for c in cols]
    for field, keywords in KEYWORD_MAP.items():
        for kw in keywords:
            for i, col in enumerate(cols_lower):
                if kw in col:
                    mapping[field] = cols[i]
                    break
            if field in mapping:
                break
    return mapping


def interactive_map(df: pd.DataFrame, template: dict | None = None) -> dict:
    print("\n─── Colunas disponíveis ───")
    for i, col in enumerate(df.columns):
        sample = df[col].dropna().head(2).tolist()
        print(f"  [{i}] {col:30s} → Ex: {sample}")

    auto = auto_map_columns(df)
    if template:
        auto.update(template)

    print(f"\n─── Mapeamento auto-detectado ───")
    for f in ["produto", "codigo", "dosagem", "preco"]:
        print(f"  {f:10s} → {auto.get(f, '(não detectado)')}")

    print("\nConfirmar? [Enter = sim | digitar índice para alterar]")
    for field in ["produto", "codigo", "dosagem", "preco"]:
        cur = auto.get(field, "")
        hint = "(obrigatório)" if field in ["produto", "preco"] else "(opcional, Enter para ignorar)"
        resp = input(f"  {field:10s} [{cur}] {hint}: ").strip()
        if resp:
            try:
                auto[field] = df.columns[int(resp)]
            except:
                auto[field] = resp
    return auto


def apply_mapping(df: pd.DataFrame, mapping: dict) -> pd.DataFrame:
    result = pd.DataFrame()
    if "produto" in mapping and mapping["produto"] in df.columns:
        result["Produto"] = df[mapping["produto"]].astype(str).str.strip()
    else:
        raise ValueError("Coluna 'produto' não mapeada.")
    if "preco" in mapping and mapping["preco"] in df.columns:
        result["Preco_USD"] = pd.to_numeric(
            df[mapping["preco"]].astype(str).str.replace(r"[^\d.]", "", regex=True), errors="coerce"
        )
    else:
        raise ValueError("Coluna 'preco' não mapeada.")
    if "codigo" in mapping and mapping.get("codigo") in df.columns:
        result["Codigo"] = df[mapping["codigo"]].astype(str).str.strip()
    else:
        result["Codigo"] = ""
    if "dosagem" in mapping and mapping.get("dosagem") in df.columns:
        result["Dosagem"] = df[mapping["dosagem"]].astype(str).str.strip()
    else:
        result["Dosagem"] = ""
    result = result[result["Produto"].str.len() > 1]
    result = result[result["Preco_USD"] > 0]
    result.dropna(subset=["Preco_USD"], inplace=True)
    return result.reset_index(drop=True)


# ─── TEMPLATES ───────────────────────────────────────────────────────
def load_templates() -> dict:
    if Path(TEMPLATES_FILE).exists():
        return json.loads(Path(TEMPLATES_FILE).read_text(encoding="utf-8"))
    return {}


def save_template(supplier: str, mapping: dict):
    templates = load_templates()
    templates[supplier] = mapping
    Path(TEMPLATES_FILE).write_text(json.dumps(templates, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  💾 Template salvo para '{supplier}'")


# ─── MAIN ────────────────────────────────────────────────────────────
def process_file(path: Path, supplier: str = "", use_ocr: bool = False,
                 save_tmpl: bool = False, silent: bool = False) -> Path | None:
    ext = path.suffix.lower().lstrip(".")
    print(f"\n{'='*55}")
    print(f"  Arquivo : {path.name}")
    print(f"  Formato : {ext.upper()}")

    try:
        if ext == "csv":
            df = extract_csv(path)
        elif ext in ["xlsx", "xls"]:
            df = extract_xlsx(path)
        elif ext == "pdf":
            if use_ocr:
                df = extract_pdf_ocr(path)
            else:
                try:
                    df = extract_pdf_text(path)
                except Exception as e:
                    print(f"  ⚠  Texto falhou ({e}), tentando OCR...")
                    df = extract_pdf_ocr(path)
        elif ext in ["jpg", "jpeg", "png"]:
            df = extract_image_ocr(path)
        elif ext in ["docx", "doc"]:
            df = extract_docx(path)
        else:
            print(f"  ✗ Formato .{ext} não suportado")
            return None
    except Exception as e:
        print(f"  ✗ Erro na extração: {e}")
        return None

    # Mapeamento
    templates = load_templates()
    tmpl = templates.get(supplier) if supplier else None
    auto = auto_map_columns(df)
    if tmpl:
        auto.update(tmpl)

    if not silent:
        mapping = interactive_map(df, auto)
    else:
        mapping = auto
        print(f"  ℹ  Mapeamento automático: {mapping}")

    try:
        result = apply_mapping(df, mapping)
    except ValueError as e:
        print(f"  ✗ Mapeamento falhou: {e}")
        return None

    # Adicionar fornecedor
    if supplier:
        result.insert(0, "Fornecedor", supplier)

    # Salvar CSV
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_name = f"{path.stem}_{supplier or 'extraido'}_aura.csv"
    out_path = Path(OUTPUT_DIR) / out_name
    result.to_csv(out_path, index=False, encoding="utf-8")
    print(f"\n  ✅ Salvo: {out_path}")
    print(f"     {len(result)} produtos com preços válidos")

    if save_tmpl and supplier:
        save_template(supplier, mapping)

    return out_path


def process_folder(folder: Path, **kwargs):
    exts = {".csv", ".xlsx", ".xls", ".pdf", ".docx", ".jpg", ".jpeg", ".png"}
    files = [f for f in folder.iterdir() if f.suffix.lower() in exts]
    print(f"\n📂 {len(files)} arquivo(s) encontrados em {folder}")
    for f in files:
        process_file(f, **kwargs)


def main():
    parser = argparse.ArgumentParser(description="AURA ERP — Extrator Universal")
    parser.add_argument("input", help="Arquivo ou pasta a processar")
    parser.add_argument("--fornecedor", "-f", default="", help="Nome do fornecedor")
    parser.add_argument("--ocr", action="store_true", help="Forçar OCR (PDFs escaneados)")
    parser.add_argument("--template", "-t", action="store_true", help="Salvar mapeamento como template")
    parser.add_argument("--auto", "-a", action="store_true", help="Modo automático (sem perguntas)")
    args = parser.parse_args()

    path = Path(args.input)
    if not path.exists():
        sys.exit(f"✗ Arquivo/pasta não encontrado: {path}")

    if path.is_dir():
        process_folder(path, supplier=args.fornecedor, use_ocr=args.ocr,
                       save_tmpl=args.template, silent=args.auto)
    else:
        process_file(path, supplier=args.fornecedor, use_ocr=args.ocr,
                     save_tmpl=args.template, silent=args.auto)

    print(f"\n📥 CSVs prontos para importar em: {OUTPUT_DIR}/")
    print("   Abra o AURA ERP → aba Importar → arraste o CSV gerado")


if __name__ == "__main__":
    main()
